"""
Заполнение .docx шаблона данными КП:
- подстановка меток {ОРГАНИЗАЦИЯ}, {ИТОГО}, {НОМЕР_ДОКУМЕНТА} и т.п.
- замена маркера {ТАБЛИЦА_ПОЗИЦИЙ} на полноценную Word-таблицу с позициями
Возвращает bytes готового .docx (для последующей конвертации в PDF).
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def format_money(val):
    try:
        f = float(val)
        return '{:,.2f}'.format(f).replace(',', '\u00a0').replace('.', ',')
    except Exception:
        return str(val)


def _replace_in_paragraph(paragraph, replacements):
    """Заменяет метки в параграфе, сохраняя форматирование первого run."""
    full = ''.join(run.text for run in paragraph.runs)
    if '{' not in full:
        return
    new = full
    for key, val in replacements.items():
        new = new.replace('{' + key + '}', str(val))
    if new == full:
        return
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new)


def _shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def _set_cell_text(cell, text, bold=False, color=None, align='left', size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _build_items_table(doc, anchor_paragraph, items, total):
    """Создаёт Word-таблицу позиций и вставляет её на место anchor_paragraph."""
    headers = ['№', 'Наименование услуги', 'Ед.изм.', 'Кол-во', 'Цена, руб.', 'Сумма, руб.']
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.autofit = True

    # Заголовок
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_text(hdr[i], h, bold=True, color=(255, 255, 255),
                       align='center' if i != 1 else 'left', size=9)
        _shade_cell(hdr[i], '1A0A6E')

    # Строки
    for idx, item in enumerate(items, 1):
        qty = item.get('qty', 1)
        price = item.get('price', 0)
        it_total = item.get('total', float(qty) * float(price))
        cells = table.add_row().cells
        _set_cell_text(cells[0], idx, align='center', size=9)
        _set_cell_text(cells[1], item.get('name', ''), align='left', size=9)
        _set_cell_text(cells[2], item.get('unit', 'шт.'), align='center', size=9)
        _set_cell_text(cells[3], qty, align='center', size=9)
        _set_cell_text(cells[4], format_money(price), align='center', size=9)
        _set_cell_text(cells[5], format_money(it_total), align='center', size=9)
        if idx % 2 == 0:
            for c in cells:
                _shade_cell(c, 'F2F2F7')

    # Итог
    total_cells = table.add_row().cells
    total_cells[0].merge(total_cells[4])
    _set_cell_text(total_cells[0], 'ИТОГО:', bold=True, align='right', size=10)
    _set_cell_text(total_cells[5], format_money(total) + ' руб.', bold=True, align='center', size=10)
    for c in (total_cells[0], total_cells[5]):
        _shade_cell(c, 'E8E8EF')

    # Перемещаем таблицу на место маркера
    anchor_paragraph._p.addprevious(table._tbl)
    # Удаляем маркерный параграф
    anchor_paragraph._p.getparent().remove(anchor_paragraph._p)


def fill_docx(template_bytes, replacements, items, total):
    """Заполняет шаблон и возвращает bytes готового .docx."""
    doc = Document(io.BytesIO(template_bytes))
    TABLE_MARKER = '{ТАБЛИЦА_ПОЗИЦИЙ}'

    # Заменяем метки в обычных параграфах + ищем маркер таблицы
    table_anchor = None
    for paragraph in doc.paragraphs:
        text = ''.join(r.text for r in paragraph.runs)
        if TABLE_MARKER in text:
            table_anchor = paragraph
        else:
            _replace_in_paragraph(paragraph, replacements)

    # Заменяем метки в существующих таблицах шаблона
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, replacements)

    # Вставляем таблицу позиций
    if table_anchor is not None:
        _build_items_table(doc, table_anchor, items, total)
    else:
        # Если маркера нет — добавляем таблицу в конец документа
        doc.add_paragraph()
        anchor = doc.add_paragraph()
        _build_items_table(doc, anchor, items, total)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
