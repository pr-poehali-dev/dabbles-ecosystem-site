# v4
import json
import os
import io
import base64
import uuid
from datetime import datetime
import psycopg2
import boto3
from docx import Document
from pdf_builder import build_kp_pdf

CORS = {
    'Access-Control-Allow-Origin': '*',
    'Access-Control-Allow-Methods': 'GET, POST, OPTIONS',
    'Access-Control-Allow-Headers': 'Content-Type, X-Auth-Token',
}

def db():
    return psycopg2.connect(os.environ['DATABASE_URL'])

def s3_client():
    return boto3.client(
        's3',
        endpoint_url='https://bucket.poehali.dev',
        aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
        aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY'],
    )

def resp(status, body):
    return {
        'statusCode': status,
        'headers': {**CORS, 'Content-Type': 'application/json'},
        'isBase64Encoded': False,
        'body': json.dumps(body, ensure_ascii=False, default=str),
    }

def esc(v):
    if v is None:
        return 'NULL'
    if isinstance(v, bool):
        return 'TRUE' if v else 'FALSE'
    if isinstance(v, (int, float)):
        return str(v)
    return "'" + str(v).replace("'", "''") + "'"

def get_admin(conn, token, schema):
    if not token:
        return None
    safe = token.replace("'", "''")
    with conn.cursor() as cur:
        cur.execute(
            f"SELECT u.id, u.role FROM {schema}.sessions s JOIN {schema}.users u ON s.user_id = u.id "
            f"WHERE s.token = '{safe}' AND s.expires_at > NOW() AND u.is_active = TRUE"
        )
        row = cur.fetchone()
    if not row or row[1] != 'admin':
        return None
    return {'id': row[0]}

def format_money(val):
    try:
        f = float(val)
        return '{:,.2f}'.format(f).replace(',', ' ')
    except Exception:
        return str(val)

def check_stopwords(text, stopwords):
    text_lower = text.lower()
    for sw in stopwords:
        if sw.lower() in text_lower:
            return sw
    return None

def apply_replacements(text, replacements):
    for key, val in replacements.items():
        text = text.replace('{' + key + '}', str(val))
    return text

def extract_docx_blocks(doc, replacements):
    """Читает docx, возвращает список блоков для PDF."""
    TABLE_MARKER = '{ТАБЛИЦА_ПОЗИЦИЙ}'
    blocks = []

    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            para = para_map.get(child)
            if para is None:
                continue
            raw = ''.join(r.text for r in para.runs)
            txt = apply_replacements(raw, replacements)
            bold = any(r.bold for r in para.runs if r.text.strip())
            align_str = str(para.alignment)
            align = 1 if 'CENTER' in align_str else (2 if 'RIGHT' in align_str else 0)
            if TABLE_MARKER in txt:
                blocks.append({'type': 'table_marker'})
            else:
                blocks.append({'type': 'para', 'text': txt, 'bold': bold, 'align': align})
        elif tag == 'tbl':
            tbl = table_map.get(child)
            if tbl is None:
                continue
            has_marker = any(
                TABLE_MARKER in apply_replacements(''.join(p.text for p in cell.paragraphs), replacements)
                for row in tbl.rows for cell in row.cells
            )
            if has_marker:
                blocks.append({'type': 'table_marker'})
            else:
                rows_data = [
                    [apply_replacements(''.join(p.text for p in cell.paragraphs), replacements) for cell in row.cells]
                    for row in tbl.rows
                ]
                blocks.append({'type': 'docx_table', 'rows': rows_data})

    return blocks

def handler(event, context):
    """Генерирует КП (PDF) из .docx шаблона с проверкой стоп-слов и автонумерацией."""
    method = event.get('httpMethod', 'GET')
    if method == 'OPTIONS':
        return {'statusCode': 200, 'headers': CORS, 'isBase64Encoded': False, 'body': ''}

    qs = event.get('queryStringParameters') or {}
    action = qs.get('action', '')
    headers_in = event.get('headers') or {}
    token = headers_in.get('X-Auth-Token') or headers_in.get('x-auth-token') or ''

    conn = db()
    client = s3_client()
    schema = os.environ.get('MAIN_DB_SCHEMA', 'public')

    # === GET TEMPLATE ===
    if action == 'get-template' and method == 'GET':
        with conn.cursor() as cur:
            cur.execute(f"SELECT id, name, file_url, uploaded_at FROM {schema}.kp_templates WHERE is_active=TRUE ORDER BY id DESC LIMIT 1")
            row = cur.fetchone()
        conn.close()
        if not row:
            return resp(404, {'error': 'Шаблон не загружен'})
        return resp(200, {'id': row[0], 'name': row[1], 'file_url': row[2], 'uploaded_at': row[3]})

    # === UPLOAD TEMPLATE ===
    if action == 'upload-template' and method == 'POST':
        if not get_admin(conn, token, schema):
            conn.close()
            return resp(403, {'error': 'Доступ запрещён'})
        body = json.loads(event.get('body') or '{}')
        file_b64 = body.get('file_base64', '')
        file_name = body.get('file_name', 'template.docx')
        if not file_b64:
            conn.close()
            return resp(400, {'error': 'Файл не передан'})
        file_bytes = base64.b64decode(file_b64)
        key = f'kp-templates/{uuid.uuid4()}/{file_name}'
        client.put_object(Bucket='files', Key=key, Body=file_bytes,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        cdn_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{key}"
        with conn.cursor() as cur:
            cur.execute(f"UPDATE {schema}.kp_templates SET is_active=FALSE")
            cur.execute(f"INSERT INTO {schema}.kp_templates (name,file_url,is_active) VALUES ({esc(file_name)},{esc(cdn_url)},TRUE)")
        conn.commit()
        conn.close()
        return resp(200, {'ok': True, 'url': cdn_url})

    # === STOPWORDS ===
    if action == 'stopwords':
        if method == 'GET':
            is_admin = get_admin(conn, token, schema)
            with conn.cursor() as cur:
                cur.execute(f"SELECT id, word, created_at FROM {schema}.kp_stopwords ORDER BY id DESC")
                rows = cur.fetchall()
            conn.close()
            words = [{'id': r[0], 'word': r[1], 'created_at': r[2]} for r in rows]
            if is_admin:
                return resp(200, {'words': words})
            return resp(200, {'words': [w['word'] for w in words]})
        if method == 'POST':
            if not get_admin(conn, token, schema):
                conn.close()
                return resp(403, {'error': 'Доступ запрещён'})
            body = json.loads(event.get('body') or '{}')
            word = (body.get('word') or '').strip()
            if not word:
                conn.close()
                return resp(400, {'error': 'Слово не передано'})
            with conn.cursor() as cur:
                cur.execute(f"INSERT INTO {schema}.kp_stopwords (word) VALUES ({esc(word)}) ON CONFLICT (word) DO NOTHING")
            conn.commit()
            conn.close()
            return resp(200, {'ok': True})
        if method == 'DELETE':
            if not get_admin(conn, token, schema):
                conn.close()
                return resp(403, {'error': 'Доступ запрещён'})
            body = json.loads(event.get('body') or '{}')
            wid = int(body.get('id') or 0)
            if not wid:
                conn.close()
                return resp(400, {'error': 'id обязателен'})
            with conn.cursor() as cur:
                cur.execute(f"DELETE FROM {schema}.kp_stopwords WHERE id={wid}")
            conn.commit()
            conn.close()
            return resp(200, {'ok': True})

    # === GENERATE KP ===
    if action == 'generate' and method == 'POST':
        body = json.loads(event.get('body') or '{}')
        organization = body.get('organization', '')
        director_name = body.get('director_name', '')
        items = body.get('items', [])

        if not organization or not director_name or not items:
            conn.close()
            return resp(400, {'error': 'Заполните все обязательные поля'})

        with conn.cursor() as cur:
            cur.execute(f"SELECT word FROM {schema}.kp_stopwords")
            stopwords = [r[0] for r in cur.fetchall()]

        check_text = organization + ' ' + director_name + ' ' + ' '.join(it.get('name','') for it in items)
        if check_stopwords(check_text, stopwords):
            conn.close()
            return resp(403, {'error': 'DENIED', 'code': 'stopword_match'})

        with conn.cursor() as cur:
            cur.execute(f"SELECT file_url FROM {schema}.kp_templates WHERE is_active=TRUE ORDER BY id DESC LIMIT 1")
            row = cur.fetchone()
        if not row:
            conn.close()
            return resp(404, {'error': 'Шаблон КП не загружен. Обратитесь к администратору.'})

        template_url = row[0]

        with conn.cursor() as cur:
            cur.execute(f"SELECT nextval('{schema}.kp_doc_seq')")
            seq_num = cur.fetchone()[0]
        doc_number = f"89-101/{datetime.now().year}-{seq_num}"

        prefix = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/"
        s3_key = template_url.replace(prefix, '')
        obj = client.get_object(Bucket='files', Key=s3_key)
        template_bytes = obj['Body'].read()

        total = sum(float(it.get('total', 0) or 0) for it in items)

        replacements = {
            'ОРГАНИЗАЦИЯ': organization,
            'organization': organization,
            'ФИО_руководителя': director_name,
            'director_name': director_name,
            'ДАТА': datetime.now().strftime('%d.%m.%Y'),
            'date': datetime.now().strftime('%d.%m.%Y'),
            'ИТОГО': format_money(total) + ' руб.',
            'total': format_money(total),
            'НОМЕР_ДОКУМЕНТА': doc_number,
            'doc_number': doc_number,
            'НОМ': doc_number,
        }

        doc_obj = Document(io.BytesIO(template_bytes))
        blocks = extract_docx_blocks(doc_obj, replacements)
        pdf_bytes = build_kp_pdf(blocks, items, total,
                                 organization=organization,
                                 director_name=director_name,
                                 doc_number=doc_number)

        safe_org = organization[:30].replace('/', '-').replace('"', '')
        result_key = f'kp-results/{uuid.uuid4()}/KP_{safe_org}.pdf'
        client.put_object(
            Bucket='files', Key=result_key, Body=pdf_bytes,
            ContentType='application/pdf',
            ContentDisposition='attachment; filename="KP.pdf"'
        )
        result_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{result_key}"

        with conn.cursor() as cur:
            items_json = json.dumps(items, ensure_ascii=False)
            cur.execute(
                f"INSERT INTO {schema}.kp_requests (organization,director_name,items,total_amount,result_url,doc_number) "
                f"VALUES ({esc(organization)},{esc(director_name)},{esc(items_json)}::jsonb,{esc(total)},{esc(result_url)},{esc(doc_number)})"
            )
        conn.commit()
        conn.close()
        return resp(200, {'ok': True, 'download_url': result_url, 'doc_number': doc_number})

    conn.close()
    return resp(404, {'error': 'Неизвестное действие'})